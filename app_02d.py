import streamlit as st
import pandas as pd
from docx import Document
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import tempfile
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import json

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("GOOGLE_API_KEY not found in environment variables.")
    st.stop()
genai.configure(api_key=api_key)

# Hàm tải lịch sử chat từ file
def load_chat_history():
    try:
        with open("rag_chat_history.json", "r", encoding="utf-8") as file:
            return json.load(file)
    except FileNotFoundError:
        return []

# Hàm lưu lịch sử chat vào file
def save_chat_history(history):
    with open("rag_chat_history.json", "w", encoding="utf-8") as file:
        json.dump(history, file, ensure_ascii=False, indent=2)

def get_text_from_documents(docs):
    """Extract text from uploaded documents (PDF, TXT, XLSX, CSV, DOCX)."""
    text = ""
    try:
        for file in docs:
            suffix = os.path.splitext(file.name)[-1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(file.read())
                tmp_path = tmp_file.name

            # Loaders theo loại file
            if suffix == ".pdf":
                loader = PyPDFLoader(tmp_path)
                pages = loader.load_and_split()
                text += "\n".join([p.page_content for p in pages])

            elif suffix == ".txt":
                loader = TextLoader(tmp_path)
                pages = loader.load_and_split()
                text += "\n".join([p.page_content for p in pages])

            elif suffix == ".csv":
                df = pd.read_csv(tmp_path)
                text += df.to_string(index=False)

            elif suffix == ".xlsx":
                try:
                    sheets = pd.read_excel(tmp_path, sheet_name=None)
                    for name, df in sheets.items():
                        text += f"\n--- Sheet: {name} ---\n"
                        text += df.to_string(index=False)
                except Exception as e:
                    st.warning(f"Lỗi đọc Excel: {e}")
                    continue
                
            elif suffix == ".docx":
                try:
                    doc = Document(tmp_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                        text += "\n" # Thêm xuống dòng sau mỗi bảng
                        
                    text += "\n" # Thêm xuống dòng sau mỗi file docx để tách biệt
                except Exception as e:
                    st.warning(f"Lỗi đọc file DOCX: {e}")
                    continue
                
            else:
                st.warning(f"Không hỗ trợ định dạng file: {suffix}")
                continue

            os.unlink(tmp_path)

    except Exception as e:
        st.error(f"Lỗi xử lý tài liệu: {str(e)}")
        return ""
    return text

def get_text_chunks(text):
    """Split text into chunks."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=800)
        chunks = text_splitter.split_text(text)
        return chunks
    except Exception as e:
        st.error(f"Error splitting text: {str(e)}")
        return []

def get_vector_store(text_chunks):
    """Create and save FAISS vector store."""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        st.success("Tài liệu đã phân tích xong, sẵn sàng để trả lời câu hỏi.")
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")

def get_prompt_template():
    """Returns the prompt template."""
    prompt_template = """
    Bạn là một chuyên viên phân tích báo cáo tài chính chuyên nghiệp, có nhiệm vụ hỗ trợ người dùng phân tích dữ liệu tài chính doanh nghiệp theo từng năm. Dữ liệu thường được trích xuất từ bảng cân đối kế toán, bao gồm các chỉ tiêu sau:

    - Tài sản ngắn hạn (TSNH)
    - Tổng tài sản (TTS)
    - Tổng nợ phải trả (Tổng Nợ)
    - Nợ ngắn hạn (Nợ NH)
    - Vốn chủ sở hữu (VCSH)

    Nhiệm vụ của bạn:

    - Phân tích các số liệu tài chính do người dùng cung cấp, kể cả theo từng năm nếu có.
    - Tính các chỉ số tài chính cơ bản và chuyên sâu, ưu tiên sử dụng công thức do người dùng cung cấp nếu có. Nếu không, hãy sử dụng các công thức mặc định sau:
        - Tỷ suất tài sản ngắn hạn = TSNH / TTS
        - Tỷ lệ nợ = Tổng Nợ / TTS
        - Tỷ lệ vốn chủ sở hữu = VCSH / TTS
        - Đòn bẩy tài chính (Hệ số Tài sản trên Vốn chủ sở hữu) = TTS / VCSH
        - Khả năng thanh toán hiện hành = TSNH / Nợ NH
        - Vốn lưu động thuần = TSNH - Nợ NH
        - Tăng trưởng tài sản = (TTS năm sau – TTS năm trước) / TTS năm trước
        - Tỷ lệ Nợ trên Vốn chủ sở hữu = Tổng Nợ / VCSH
        - Tài sản dài hạn (hoặc Tài sản cố định) = TTS - TSNH
        - Nợ dài hạn = Tổng Nợ - Nợ NH
        - (và các chỉ số khác có thể suy ra từ dữ liệu đã cho)
    - Đánh giá ý nghĩa tài chính của các số liệu và chỉ số đã tính toán, cung cấp nhận định chuyên nghiệp và bối cảnh nếu có thể (ví dụ: so sánh với các năm trước, ý nghĩa tích cực/tiêu cực).

    Ví dụ 1:
    User: Trong năm 2024, tài sản ngắn hạn là 37.553 tỷ, tổng tài sản là 55.049 tỷ. Hãy tính tỷ suất tài sản ngắn hạn và cho tôi biết doanh nghiệp có ổn định không?

    Assistant:
    Tỷ suất tài sản ngắn hạn năm 2024 là 37.553 / 55.049 ≈ 68,22%. Đây là một tỷ lệ cao, cho thấy doanh nghiệp có khả năng thanh toán ngắn hạn rất tốt. Nếu so sánh với các năm trước mà vẫn giữ ổn định quanh mức này thì đây là một dấu hiệu tài chính tích cực.

    Ví dụ 2:
    User: Tổng nợ năm 2024 là 19.500 tỷ, tổng tài sản là 55.000 tỷ. Hãy phân tích tỷ lệ nợ giúp tôi.

    Assistant:
    Tỷ lệ nợ năm 2024 = 19.500 / 55.000 = 35,45%. Tỷ lệ này ở mức hợp lý. Nếu phần còn lại được tài trợ bởi vốn chủ sở hữu thì doanh nghiệp đang vận hành với mức đòn bẩy an toàn.

    Ví dụ 3:
    User: Năm 2023, tài sản ngắn hạn là 35.900 tỷ, nợ ngắn hạn là 10.000 tỷ. Tôi muốn biết vốn lưu động thuần.

    Assistant:
    Vốn lưu động thuần = 35.900 – 10.000 = 25.900 tỷ. Đây là phần dư tài sản ngắn hạn sau khi trừ đi nợ ngắn hạn, thể hiện khả năng duy trì hoạt động và dự phòng ngắn hạn của doanh nghiệp.

    Ví dụ 4:
    User: Tổng tài sản năm 2023 là 52.600 tỷ, năm 2024 là 55.000 tỷ. Hãy tính tăng trưởng tổng tài sản.

    Assistant:
    Tăng trưởng tổng tài sản = (55.000 – 52.600) / 52.600 = 4,57%. Doanh nghiệp có mức tăng trưởng tổng tài sản dương, cho thấy quy mô tài sản tiếp tục được mở rộng.

    Ví dụ 5 (Mới - Minh họa ưu tiên công thức người dùng):
    User: Tài sản ngắn hạn năm 2023 là 30.000 tỷ, tổng tài sản năm 2023 là 50.000 tỷ. Hãy tính Tỷ suất tài sản ngắn hạn của năm 2023, biết rằng Tỷ suất tài sản ngắn hạn = (tài sản ngắn hạn / tổng tài sản) * 100%.

    Assistant:
    Theo công thức bạn cung cấp, Tỷ suất tài sản ngắn hạn năm 2023 = (30.000 / 50.000) * 100% = 60%. Đây là một tỷ lệ cao, cho thấy doanh nghiệp có cơ cấu tài sản ưu tiên yếu tố thanh khoản ngắn hạn.

    Ghi chú:
    Nếu dữ liệu người dùng chưa đầy đủ để tính, hãy hỏi lại thông tin còn thiếu một cách ngắn gọn và gợi ý đúng mã số cần cung cấp.

    Ngữ cảnh: {context}
    Câu hỏi: {question}

    Answer:
    """
    try:
        return PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    except Exception as e:
        st.error(f"Error creating prompt template: {str(e)}")
        return None

def user_input(user_question):
    """Process user question and return streamed response, integrated with chat history.""" 
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        
        if not os.path.exists("faiss_index"):
            st.error("Không tìm thấy FAISS index. Hãy tải tài liệu lên trước.")
            return
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)

        prompt_template_obj = get_prompt_template() # Lấy prompt template
        if not prompt_template_obj:
            return

        # Nối nội dung các tài liệu tìm được vào ngữ cảnh
        context = "\n\n".join([doc.page_content for doc in docs])

        # Tạo prompt cuối cùng đã bổ sung ngữ cảnh RAG
        rag_augmented_question = prompt_template_obj.format(context=context, question=user_question)
        
        # Chuyển đổi lịch sử chat của Streamlit sang định dạng của Google Generative AI
        # Loại bỏ tin nhắn hiện tại của người dùng khỏi lịch sử để tránh trùng lặp khi gửi message
        gemini_history = []
        for msg in st.session_state.chat_history:
            if msg["role"] == "user":
                gemini_history.append({"role": "user", "parts": [msg["content"]]})
            elif msg["role"] == "assistant":
                gemini_history.append({"role": "model", "parts": [msg["content"]]})

        # Khởi tạo model và chat session với lịch sử
        model = genai.GenerativeModel("gemini-2.0-flash")
        chat = model.start_chat(history=gemini_history)

        st.write("Reply:") 
        response_placeholder = st.empty()
        full_response_text = ""

        # Stream phản hồi từ mô hình sử dụng chat.send_message
        for chunk in chat.send_message(rag_augmented_question, stream=True):
            if chunk.text:
                full_response_text += chunk.text
                response_placeholder.markdown(full_response_text)
        
        # Lưu phản hồi của bot vào session state và file
        st.session_state.chat_history.append({"role": "assistant", "content": full_response_text})
        save_chat_history(st.session_state.chat_history)

    except Exception as e:
        st.error(f"Lỗi xử lý câu hỏi: {str(e)}")

def main():
    st.set_page_config(page_title="Read Documents", page_icon="📄")
    st.header("Demo phân tích tài liệu 📄")

    # Khởi tạo và tải lịch sử chat
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = load_chat_history()

    # Hiển thị lịch sử chat
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    user_question = st.chat_input("Bạn hãy hỏi sau khi tài liệu đã được phân tích xong")

    if user_question:
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.markdown(user_question)

        user_input(user_question)

    with st.sidebar:
        st.title("Menu")
        pdf_docs = st.file_uploader(
            "Tải tài liệu lên (PDF, TXT, XLSX, CSV, DOCX)",
            accept_multiple_files=True,
            type=["pdf", "docx", "txt", "xlsx", "csv"]
        )
        if st.button("Phân tích tài liệu"):
            if not pdf_docs:
                st.error("Vui lòng tải tài liệu của bạn lên.")
                return
            with st.spinner("Đang xử lý..."):
                raw_text = get_text_from_documents(pdf_docs)
                if raw_text:
                    text_chunks = get_text_chunks(raw_text)
                    if text_chunks:
                        get_vector_store(text_chunks)
                    else:
                        st.error("Kiểm tra lại nội dung trong tài liệu.")
                else:
                    st.error("Không có nội dung nào được trong tài liệu.")
        
        # Nút xóa lịch sử chat
        if st.button("Xóa lịch sử chat"):
            st.session_state.chat_history = []
            if os.path.exists("rag_chat_history.json"):
                os.remove("rag_chat_history.json")
            st.rerun()

if __name__ == "__main__":
    main()