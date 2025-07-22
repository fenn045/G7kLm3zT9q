import streamlit as st
import pandas as pd
from docx import Document
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import tempfile
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("GOOGLE_API_KEY not found in environment variables.")
    st.stop()
genai.configure(api_key=api_key)

def get_text_from_documents(docs):
    """Extract text from uploaded documents (PDF, TXT, XLSX, CSV, DOCX)."""
    text = ""
    try:
        for file in docs:
            suffix = os.path.splitext(file.name)[-1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(file.read())
                tmp_path = tmp_file.name

            # Loaders theo loại file
            if suffix == ".pdf":
                loader = PyPDFLoader(tmp_path)
                pages = loader.load_and_split()
                text += "\n".join([p.page_content for p in pages])

            elif suffix == ".txt":
                loader = TextLoader(tmp_path)
                pages = loader.load_and_split()
                text += "\n".join([p.page_content for p in pages])

            elif suffix == ".csv":
                df = pd.read_csv(tmp_path)
                text += df.to_string(index=False)

            elif suffix == ".xlsx":
                try:
                    sheets = pd.read_excel(tmp_path, sheet_name=None)
                    for name, df in sheets.items():
                        text += f"\n--- Sheet: {name} ---\n"
                        text += df.to_string(index=False)
                except Exception as e:
                    st.warning(f"Lỗi đọc Excel: {e}")
                    continue
                
            elif suffix == ".docx":
                try:
                    doc = Document(tmp_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    # Có thể bổ sung đọc text từ bảng (tables) nếu cần #
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                        text += "\n" # Thêm xuống dòng sau mỗi bảng #
                    text += "\n" # Thêm xuống dòng sau mỗi file docx để tách biệt #
                except Exception as e:
                    st.warning(f"Lỗi đọc file DOCX: {e}")
                    continue
                
            else:
                st.warning(f"Không hỗ trợ định dạng file: {suffix}")
                continue

            os.unlink(tmp_path)

    except Exception as e:
        st.error(f"Lỗi xử lý tài liệu: {str(e)}")
        return ""
    return text

def get_text_chunks(text):
    """Split text into chunks."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
        chunks = text_splitter.split_text(text)
        return chunks
    except Exception as e:
        st.error(f"Error splitting text: {str(e)}")
        return []

def get_vector_store(text_chunks):
    """Create and save FAISS vector store."""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        st.success("Tài liệu đã phân tích xong, sẵn sàng để trả lời câu hỏi.")
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")

def get_conversational_chain():
    """Create a conversational QA chain."""
    prompt_template = """
    Trả lời câu hỏi một cách chi tiết nhất có thể dựa trên ngữ cảnh được cung cấp. Nếu câu trả lời không có trong ngữ cảnh được cung cấp, hãy nói, "Câu trả lời không có trong ngữ cảnh."
    Không cung cấp thông tin sai lệch.

    Ngữ cảnh: {context}
    Câu hỏi: {question}

    Answer:
    """
    try:
        model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.7)
        prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
        return chain
    except Exception as e:
        st.error(f"Error initializing QA chain: {str(e)}")
        return None

def user_input(user_question):
    """Process user question and return response."""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        # Load FAISS index with deserialization permission
        if not os.path.exists("faiss_index"):
            st.error("Không tìm thấy FAISS index. Hãy tải tài liệu PDF lên trước.")
            return
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)

        chain = get_conversational_chain()
        if not chain:
            return

        response = chain(
            {"input_documents": docs, "question": user_question},
            return_only_outputs=True
        )
        st.write("Reply: ", response["output_text"])
    except Exception as e:
        st.error(f"Lỗi xử lý câu hỏi: {str(e)}")

def main():
    st.set_page_config(page_title="Chatbot", page_icon="📄")
    st.header("Demo Chatbot phân tích tài liệu")

    user_question = st.text_input("Bạn hãy hỏi sau khi tài liệu đã được phân tích xong")

    if user_question:
        user_input(user_question)

    with st.sidebar:
        st.title("Menu")
        pdf_docs = st.file_uploader(
            "Tải tài liệu lên (PDF, TXT, XLSX, CSV, DOCX)",
            accept_multiple_files=True,
            type=["pdf", "docx", "txt", "xlsx", "csv"]
        )
        if st.button("Phân tích tài liệu"):
            if not pdf_docs:
                st.error("Vui lòng tải tài liệu của bạn lên.")
                return
            with st.spinner("Đang xử lý..."):
                raw_text = get_text_from_documents(pdf_docs)
                if raw_text:
                    text_chunks = get_text_chunks(raw_text)
                    if text_chunks:
                        get_vector_store(text_chunks)
                    else:
                        st.error("Kiểm tra lại nội dung trong tài liệu")
                else:
                    st.error("Không có nội dung nào được trong tài liệu.")

if __name__ == "__main__":
    main()